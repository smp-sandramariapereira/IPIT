#!/usr/bin/env python3
"""Gera o Kit IPIT em DOCX a partir dos arquivos Markdown do repositório."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "downloads"
INDIVIDUAL = OUT / "individual"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "13243A"
MUTED = "5C6877"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"

SOURCES = [
    ("docs/guia-facilitacao-docente.md", "01-Guia-de-Facilitacao-Docente.docx"),
    ("docs/plano-de-aplicacao.md", "02-Plano-de-Aplicacao.docx"),
    ("kit-gratuito/01-checklist-inicial.md", "03-Checklist-Inicial.docx"),
    ("kit-gratuito/02-canvas-problema.md", "04-Canvas-do-Problema.docx"),
    ("kit-gratuito/03-canvas-proposta-de-valor.md", "05-Canvas-da-Proposta-de-Valor.docx"),
    ("kit-gratuito/04-roteiro-pitch.md", "06-Roteiro-de-Pitch.docx"),
    ("kit-gratuito/05-rubrica-resumida.md", "07-Rubrica-Resumida.docx"),
    ("kit-gratuito/06-guia-micro-ideathon.md", "08-Guia-do-Micro-Ideathon.docx"),
    ("kit-gratuito/07-guia-uso-responsavel-ia.md", "09-Uso-Responsavel-de-IA.docx"),
]

EXAMPLE_SOURCES = [
    "exemplos/micro-ideathon-exemplo-ficticio/README.md",
    "exemplos/micro-ideathon-exemplo-ficticio/01-problema.md",
    "exemplos/micro-ideathon-exemplo-ficticio/02-proposta-de-valor.md",
    "exemplos/micro-ideathon-exemplo-ficticio/03-prototipo.md",
    "exemplos/micro-ideathon-exemplo-ficticio/04-pitch.md",
    "exemplos/micro-ideathon-exemplo-ficticio/05-rubrica-preenchida.md",
    "exemplos/micro-ideathon-exemplo-ficticio/06-retrospectiva.md",
]


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("IPIT  |  Kit do Professor")
    set_run_font(run, 8.5, MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Sandra Maria Pereira  |  CC BY 4.0  |  Página ")
    set_run_font(run, 8.5, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    doc.core_properties.title = title
    doc.core_properties.author = "Sandra Maria Pereira"
    doc.core_properties.subject = "IPIT — Ideathon Pedagógico de Inovação Tecnológica"


def add_cover(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IPIT")
    set_run_font(run, 12, BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    set_run_font(run, 28, INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(64)
    run = p.add_run(subtitle)
    set_run_font(run, 14, DARK_BLUE, italic=True)

    for text, bold in [
        ("Ideathon Pedagógico de Inovação Tecnológica", True),
        ("Autora: Sandra Maria Pereira", False),
        ("Versão 1.0.0 — setembro de 2026", False),
        ("Conteúdo educacional licenciado sob CC BY 4.0", False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_run_font(run, 10.5, MUTED, bold=bold)
    doc.add_page_break()


def add_inline(paragraph, text: str):
    def readable_link(match):
        label, target = match.group(1), match.group(2)
        return f"{label} ({target})" if target.startswith(("http://", "https://")) else label

    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", readable_link, text)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    if cols == 2:
        widths = [2700, 6660]
    if cols == 3:
        widths = [2300, 2060, 5000]
    if cols == 4:
        widths = [1800, 1200, 3180, 3180]
    if cols >= 5:
        widths = [1872] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, values in enumerate(normalized):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value)
            if r_idx == 0:
                shade_cell(cell, HEADER_FILL)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(INK)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [part.strip() for part in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part.replace(" ", "")) for part in parts):
            rows.append(parts)
        i += 1
    return rows, i


def clean_heading(text: str) -> str:
    text = re.sub(r"^[^\wÀ-ÿ]+\s*", "", text)
    return text.strip()


def create_list_numbering(doc: Document, kind: str) -> int:
    """Cria uma lista independente para que a numeração reinicie corretamente."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_list_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_pr.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(number)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_markdown(doc: Document, path: Path, include_title=True):
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    first_h1 = True
    active_bullet_id = None
    active_number_id = None
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            i += 1
            continue
        if not stripped or stripped == "---":
            active_bullet_id = None
            active_number_id = None
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            active_bullet_id = None
            active_number_id = None
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            active_bullet_id = None
            active_number_id = None
            level = len(heading.group(1))
            text = clean_heading(heading.group(2))
            if level == 1 and first_h1 and not include_title:
                first_h1 = False
                i += 1
                continue
            first_h1 = False
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, text)
            i += 1
            continue
        if stripped.startswith(">"):
            active_bullet_id = None
            active_number_id = None
            text = stripped.lstrip("> ")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), CALLOUT_FILL)
            p_pr.append(shd)
            add_inline(p, text)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        number = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if bullet:
            active_number_id = None
            content = bullet.group(1)
            if content.startswith("[ ]"):
                active_bullet_id = None
                p = doc.add_paragraph()
                add_inline(p, "☐ " + content[3:].strip())
            else:
                if active_bullet_id is None:
                    active_bullet_id = create_list_numbering(doc, "bullet")
                p = doc.add_paragraph()
                apply_list_numbering(p, active_bullet_id)
                add_inline(p, content)
            i += 1
            continue
        if number:
            active_bullet_id = None
            if active_number_id is None:
                active_number_id = create_list_numbering(doc, "number")
            p = doc.add_paragraph()
            apply_list_numbering(p, active_number_id)
            add_inline(p, number.group(1))
            i += 1
            continue
        paragraph_lines = [stripped]
        active_bullet_id = None
        active_number_id = None
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", ">", "```")) or re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+[.)]\s+", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def build_single(source: str, output_name: str):
    path = ROOT / source
    first = next((line[2:].strip() for line in path.read_text(encoding="utf-8").splitlines() if line.startswith("# ")), path.stem)
    title = clean_heading(first)
    doc = Document()
    style_document(doc, title)
    add_cover(doc, title, "Material editável para professores")
    add_markdown(doc, path, include_title=False)
    target = INDIVIDUAL / output_name
    doc.save(target)
    return target


def build_combined():
    doc = Document()
    title = "Kit IPIT do Professor"
    style_document(doc, title)
    add_cover(doc, title, "Planeje, conduza e avalie um Micro-Ideathon")
    intro = doc.add_paragraph()
    intro.style = "Heading 1"
    intro.add_run("Como usar este kit")
    intro_numbering = create_list_numbering(doc, "number")
    for text in [
        "Comece pelo plano de aplicação e pelo checklist.",
        "Use o guia de facilitação durante o planejamento e a condução.",
        "Selecione os canvases e formulários compatíveis com seu tempo.",
        "Consulte o exemplo fictício somente como referência de preenchimento.",
        "Adapte o material ao currículo, à infraestrutura e às necessidades da turma.",
    ]:
        p = doc.add_paragraph()
        apply_list_numbering(p, intro_numbering)
        add_inline(p, text)
    combined_sources = [source for source, _ in SOURCES] + EXAMPLE_SOURCES
    for idx, source in enumerate(combined_sources):
        # Os materiais principais começam em página nova. As partes do exemplo
        # fictício fluem em sequência para evitar páginas quase vazias.
        if 0 < idx <= len(SOURCES):
            doc.add_page_break()
        add_markdown(doc, ROOT / source, include_title=True)
    target = OUT / "Kit-IPIT-Professor-Editavel.docx"
    doc.save(target)
    return target


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    INDIVIDUAL.mkdir(parents=True, exist_ok=True)
    generated = [build_single(source, name) for source, name in SOURCES]
    generated.append(build_combined())
    for path in generated:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
